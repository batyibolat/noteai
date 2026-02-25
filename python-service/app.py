# from flask import Flask, request, jsonify
# from flask_cors import CORS
# import google.generativeai as genai
# import os
# from dotenv import load_dotenv
# import tempfile
# import PyPDF2
# import docx
# import logging
# import
# # Настройка логирования
# logging.basicConfig(level=logging.INFO)
# logger = logging.getLogger(__name__)

# # Загрузка переменных окружения
# load_dotenv()

# app = Flask(__name__)
# CORS(app)

# # Конфигурация Gemini API
# GEMINI_API_KEY = os.getenv('GEMINI_API_KEY')
# if not GEMINI_API_KEY:
#     logger.error("GEMINI_API_KEY not found in environment variables")
#     raise ValueError("Please set GEMINI_API_KEY environment variable")

# genai.configure(api_key=GEMINI_API_KEY)

# def extract_text_from_pdf(file_path):
#     """Извлечение текста из PDF файла"""
#     try:
#         with open(file_path, 'rb') as file:
#             reader = PyPDF2.PdfReader(file)
#             text = ""
#             for page in reader.pages:
#                 text += page.extract_text()
#             return text
#     except Exception as e:
#         logger.error(f"Error extracting text from PDF: {e}")
#         return None

# def extract_text_from_docx(file_path):
#     """Извлечение текста из DOCX файла"""
#     try:
#         doc = docx.Document(file_path)
#         text = ""
#         for paragraph in doc.paragraphs:
#             text += paragraph.text + "\n"
#         return text
#     except Exception as e:
#         logger.error(f"Error extracting text from DOCX: {e}")
#         return None

# def summarize_text(text):
#     """Создание конспекта с помощью Gemini AI"""
#     try:
#         model = genai.GenerativeModel('gemini-2.5-flash')
#         prompt = f"""
#         Пожалуйста, создай подробный и структурированный конспект следующего текста. 
#         Конспект должен быть на русском языке и включать:
#         1. Основные идеи и ключевые моменты
#         2. Важные детали и примеры
#         3. Выводы и заключения
        
#         Текст для конспектирования:
#         {text}
        
#         Конспект:
#         """
        
#         response = model.generate_content(prompt)
#         return response.text
#     except Exception as e:
#         logger.error(f"Error generating summary: {e}")
#         return None

# @app.route('/api/summarize', methods=['POST'])
# def summarize():
#     try:
#         data = request.json
#         text = data.get('text', '')
        
#         if not text:
#             return jsonify({'error': 'No text provided'}), 400
        
#         logger.info("Generating summary...")
#         summary = summarize_text(text)
        
#         if summary:
#             return jsonify({'summary': summary})
#         else:
#             return jsonify({'error': 'Failed to generate summary'}), 500
            
#     except Exception as e:
#         logger.error(f"Error in summarize endpoint: {e}")
#         return jsonify({'error': str(e)}), 500

# @app.route('/api/summarize/file', methods=['POST'])
# def summarize_file():
#     try:
#         if 'file' not in request.files:
#             return jsonify({'error': 'No file provided'}), 400
        
#         file = request.files['file']
#         if file.filename == '':
#             return jsonify({'error': 'No file selected'}), 400
        
#         # Сохраняем файл во временную директорию
#         with tempfile.NamedTemporaryFile(delete=False) as temp_file:
#             file.save(temp_file.name)
#             temp_file_path = temp_file.name
        
#         text = None
        
#         # Определяем тип файла и извлекаем текст
#         if file.filename.lower().endswith('.pdf'):
#             text = extract_text_from_pdf(temp_file_path)
#         elif file.filename.lower().endswith(('.docx', '.doc')):
#             text = extract_text_from_docx(temp_file_path)
#         elif file.filename.lower().endswith('.txt'):
#             with open(temp_file_path, 'r', encoding='utf-8') as f:
#                 text = f.read()
#         else:
#             # Удаляем временный файл
#             os.unlink(temp_file_path)
#             return jsonify({'error': 'Unsupported file format'}), 400
        
#         # Удаляем временный файл
#         os.unlink(temp_file_path)
        
#         if not text:
#             return jsonify({'error': 'Could not extract text from file'}), 400
        
#         # Генерируем конспект
#         summary = summarize_text(text)
        
#         if summary:
#             return jsonify({
#                 'summary': summary,
#                 'original_length': len(text),
#                 'summary_length': len(summary)
#             })
#         else:
#             return jsonify({'error': 'Failed to generate summary'}), 500
            
#     except Exception as e:
#         logger.error(f"Error in summarize_file endpoint: {e}")
#         return jsonify({'error': str(e)}), 500

# @app.route('/health', methods=['GET'])
# def health_check():
#     return jsonify({'status': 'healthy', 'service': 'AI Summarizer'})

# if __name__ == '__main__':
#     logger.info("Starting AI Summarizer service...")
#     app.run(host='0.0.0.0', port=5000, debug=True)



#################################################3



# from flask import Flask, request, jsonify
# from flask_cors import CORS
# import logging
# import time

# # Настройка логирования
# logging.basicConfig(level=logging.INFO)
# logger = logging.getLogger(__name__)

# app = Flask(__name__)
# CORS(app)

# @app.route('/api/summarize', methods=['POST'])
# def summarize():
#     try:
#         data = request.json
#         text = data.get('text', '')
        
#         if not text:
#             return jsonify({'error': 'No text provided'}), 400
        
#         logger.info(f"Received text for summarization: {len(text)} characters")
        
#         # Имитация обработки AI
#         time.sleep(1)  # Задержка для имитации работы AI
        
#         # Простой конспект для тестирования
#         summary = f"Конспект текста ({len(text)} символов):\n\n"
#         summary += "Основные идеи:\n"
#         summary += "1. Первая ключевая мысль из текста\n"
#         summary += "2. Вторая важная идея\n"
#         summary += "3. Основные выводы и заключения\n\n"
#         summary += "Это автоматически сгенерированный конспект с помощью AI."
        
#         logger.info("Summary generated successfully")
#         return jsonify({'summary': summary})
            
#     except Exception as e:
#         logger.error(f"Error in summarize endpoint: {e}")
#         return jsonify({'error': str(e)}), 500

# @app.route('/api/summarize/file', methods=['POST'])
# def summarize_file():
#     try:
#         if 'file' not in request.files:
#             return jsonify({'error': 'No file provided'}), 400
        
#         file = request.files['file']
#         if file.filename == '':
#             return jsonify({'error': 'No file selected'}), 400
        
#         logger.info(f"Received file: {file.filename}")
        
#         # Имитация обработки файла
#         time.sleep(2)
        
#         summary = f"Конспект файла: {file.filename}\n\n"
#         summary += "Основное содержание файла:\n"
#         summary += "- Ключевая информация из документа\n"
#         summary += "- Важные детали и факты\n"
#         summary += "- Основные выводы\n\n"
#         summary += "Файл успешно обработан AI системой."
        
#         return jsonify({
#             'summary': summary,
#             'filename': file.filename,
#             'status': 'success'
#         })
            
#     except Exception as e:
#         logger.error(f"Error in summarize_file endpoint: {e}")
#         return jsonify({'error': str(e)}), 500

# @app.route('/health', methods=['GET'])
# def health_check():
#     return jsonify({'status': 'healthy', 'service': 'AI Summarizer'})

# @app.route('/test', methods=['GET'])
# def test():
#     return jsonify({'message': 'Python service is working!'})

# if __name__ == '__main__':
#     logger.info("Starting AI Summarizer service...")
#     app.run(host='0.0.0.0', port=5000, debug=True)

###################################3





# from flask import Flask, request, jsonify
# from flask_cors import CORS
# import logging

# logging.basicConfig(level=logging.INFO)
# logger = logging.getLogger(__name__)

# app = Flask(__name__)
# CORS(app)

# @app.route('/api/summarize', methods=['POST'])
# def summarize():
#     try:
#         data = request.json
#         text = data.get('text', '')
        
#         if not text:
#             return jsonify({'error': 'No text provided'}), 400
        
#         # Простой конспект для гарантии работы
#         summary = f"Конспект текста ({len(text)} символов):\n\n"
#         summary += "Основные идеи текста:\n"
#         summary += "1. Первая ключевая мысль\n"
#         summary += "2. Вторая важная идея\n"
#         summary += "3. Основные выводы\n\n"
#         summary += "Это автоматически сгенерированный конспект."
        
#         return jsonify({'summary': summary})
            
#     except Exception as e:
#         logger.error(f"Error: {e}")
#         return jsonify({'summary': 'Конспект создан в упрощенном режиме'})

# @app.route('/health', methods=['GET'])
# def health_check():
#     return jsonify({'status': 'healthy'})

# if __name__ == '__main__':
#     app.run(host='0.0.0.0', port=5000, debug=False)











###################FINAAAAAAAALLL########################


from flask import Flask, request, jsonify
from flask_cors import CORS
import google.generativeai as genai
import os
from dotenv import load_dotenv
import tempfile
import PyPDF2
import docx
import logging

# Настройка логирования
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Загрузка переменных окружения
load_dotenv()

app = Flask(__name__)
CORS(app)

# Конфигурация Gemini API
GEMINI_API_KEY = os.getenv('GEMINI_API_KEY')
if not GEMINI_API_KEY:
    logger.error("GEMINI_API_KEY not found in environment variables")
    raise ValueError("Please set GEMINI_API_KEY environment variable")

genai.configure(api_key=GEMINI_API_KEY)

def extract_text_from_pdf(file_path):
    """Извлечение текста из PDF файла"""
    try:
        with open(file_path, 'rb') as file:
            reader = PyPDF2.PdfReader(file)
            text = ""
            for page in reader.pages:
                text += page.extract_text()
            return text
    except Exception as e:
        logger.error(f"Error extracting text from PDF: {e}")
        return None

def extract_text_from_docx(file_path):
    """Извлечение текста из DOCX файла"""
    try:
        doc = docx.Document(file_path)
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        return text
    except Exception as e:
        logger.error(f"Error extracting text from DOCX: {e}")
        return None

def summarize_text(text):
    """Создание конспекта с помощью Gemini AI"""
    try:
        model = genai.GenerativeModel('gemini-2.5-flash')
        prompt = f"""
        Пожалуйста, создай подробный и структурированный конспект следующего текста. 
        Конспект должен быть на русском языке и включать:
        1. Основные идеи и ключевые моменты
        2. Важные детали и примеры
        3. Выводы и заключения
        
        Текст для конспектирования:
        {text}
        
        Конспект:
        """
        
        response = model.generate_content(prompt)
        return response.text
    except Exception as e:
        logger.error(f"Error generating summary: {e}")
        return None

@app.route('/api/summarize', methods=['POST'])
def summarize():
    try:
        data = request.json
        text = data.get('text', '')
        
        if not text:
            return jsonify({'error': 'No text provided'}), 400
        
        logger.info("Generating summary...")
        summary = summarize_text(text)
        
        if summary:
            return jsonify({
                'summary': summary,
                'original_text': text,
                'original_length': len(text),
                'summary_length': len(summary)
            })
        else:
            return jsonify({'error': 'Failed to generate summary'}), 500
            
    except Exception as e:
        logger.error(f"Error in summarize endpoint: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/api/summarize/file', methods=['POST'])
def summarize_file():
    try:
        if 'file' not in request.files:
            return jsonify({'error': 'No file provided'}), 400
        
        file = request.files['file']
        if file.filename == '':
            return jsonify({'error': 'No file selected'}), 400
        
        # Сохраняем файл во временную директорию
        with tempfile.NamedTemporaryFile(delete=False) as temp_file:
            file.save(temp_file.name)
            temp_file_path = temp_file.name
        
        text = None
        
        # Определяем тип файла и извлекаем текст
        if file.filename.lower().endswith('.pdf'):
            text = extract_text_from_pdf(temp_file_path)
        elif file.filename.lower().endswith(('.docx', '.doc')):
            text = extract_text_from_docx(temp_file_path)
        elif file.filename.lower().endswith('.txt'):
            with open(temp_file_path, 'r', encoding='utf-8') as f:
                text = f.read()
        else:
            # Удаляем временный файл
            os.unlink(temp_file_path)
            return jsonify({'error': 'Unsupported file format'}), 400
        
        # Удаляем временный файл
        os.unlink(temp_file_path)
        
        if not text:
            return jsonify({'error': 'Could not extract text from file'}), 400
        
        # Генерируем конспект
        logger.info("Generating summary from file...")
        summary = summarize_text(text)
        
        if summary:
            return jsonify({
                'summary': summary,
                'original_text': text,  # Возвращаем оригинальный текст
                'original_length': len(text),
                'summary_length': len(summary),
                'filename': file.filename
            })
        else:
            return jsonify({'error': 'Failed to generate summary'}), 500
            
    except Exception as e:
        logger.error(f"Error in summarize_file endpoint: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/health', methods=['GET'])
def health_check():
    return jsonify({'status': 'healthy', 'service': 'AI Summarizer'})

if __name__ == '__main__':
    logger.info("Starting AI Summarizer service...")
    app.run(host='0.0.0.0', port=5000, debug=True)
